import streamlit as st
import io
import re

# Set page configuration first
st.set_page_config(
    page_title="PDF Exam Generator",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Show loading message
with st.spinner("Loading required libraries... This may take a minute on first run."):
    try:
        import PyPDF2
        import nltk
        from transformers import pipeline
        from sentence_transformers import SentenceTransformer
        import torch
        from docx import Document
        import gc
    except ImportError as e:
        st.error(f"Error loading required libraries: {str(e)}")
        st.info("Please try refreshing the page. If the error persists, contact support.")
        st.stop()

# Initialize session state
if 'initialized' not in st.session_state:
    st.session_state.initialized = False
    st.session_state.model_loaded = False

# Cache the model loading with memory optimization
@st.cache_resource(show_spinner=False)
def load_models():
    try:
        # Load question generation model with simpler pipeline
        qa_pipeline = pipeline(
            'text2text-generation',
            model='iarfmoose/t5-base-question-generator',
            device=-1  # Force CPU
        )
        
        # Load sentence transformer with memory optimization
        sentence_model = SentenceTransformer(
            'all-MiniLM-L6-v2',
            device='cpu'
        )
        
        # Clear GPU memory if available
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
        
        # Run garbage collection
        gc.collect()
        
        return qa_pipeline, sentence_model
    except Exception as e:
        st.error(f"Error loading AI models: {str(e)}")
        return None, None

# Download required NLTK data
@st.cache_resource(show_spinner=False)
def download_nltk_data():
    try:
        nltk.data.find('tokenizers/punkt')
    except LookupError:
        with st.spinner("Downloading required language data..."):
            nltk.download('punkt', quiet=True)

# Custom CSS with loading animation
st.markdown("""
<style>
    .main {
        padding: 2rem;
    }
    .stButton > button {
        width: 100%;
        background-color: #4CAF50;
        color: white;
        padding: 0.5rem;
        margin-top: 1rem;
    }
    .stButton > button:hover {
        background-color: #45a049;
    }
    .success-message {
        padding: 1rem;
        background-color: #dff0d8;
        border-color: #d6e9c6;
        color: #3c763d;
        border-radius: 4px;
        margin: 1rem 0;
    }
    .error-message {
        padding: 1rem;
        background-color: #f2dede;
        border-color: #ebccd1;
        color: #a94442;
        border-radius: 4px;
        margin: 1rem 0;
    }
    .title-container {
        text-align: center;
        padding: 2rem 0;
        background: linear-gradient(90deg, #1e3c72 0%, #2a5298 100%);
        color: white;
        border-radius: 10px;
        margin-bottom: 2rem;
    }
    .question-box {
        background-color: #f8f9fa;
        padding: 1rem;
        border-radius: 5px;
        margin: 1rem 0;
        border-left: 5px solid #4CAF50;
    }
    .loading {
        text-align: center;
        padding: 2rem;
    }
    .footer {
        position: fixed;
        bottom: 0;
        left: 0;
        right: 0;
        background: white;
        padding: 1rem;
        text-align: center;
        border-top: 1px solid #eee;
    }
    .instructions {
        background: #f8f9fa;
        padding: 1.5rem;
        border-radius: 10px;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

class ExamGenerator:
    def __init__(self):
        # Load models with error handling
        qa_pipeline, sentence_model = load_models()
        if qa_pipeline is None or sentence_model is None:
            st.error("Failed to initialize AI models. Please try refreshing the page.")
            st.stop()
        self.qa_pipeline = qa_pipeline
        self.sentence_model = sentence_model
        
    def extract_text_from_pdf(self, pdf_file):
        try:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
            return text
        except Exception as e:
            st.error(f"Error reading PDF: {str(e)}")
            return None
    
    def preprocess_text(self, text):
        if not text:
            return []
        # Clean the text
        text = re.sub(r'\s+', ' ', text)
        text = text.strip()
        
        # Split into sentences
        sentences = nltk.sent_tokenize(text)
        
        # Filter out short sentences and those without important information
        sentences = [s for s in sentences if len(s.split()) > 5]
        return sentences
    
    def generate_questions(self, sentences, num_questions=5):
        questions = []
        
        # Generate embeddings for all sentences
        embeddings = self.sentence_model.encode(sentences)
        
        # Select diverse sentences based on embeddings
        selected_indices = []
        for _ in range(min(num_questions, len(sentences))):
            if not selected_indices:
                selected_indices.append(0)
            else:
                # Find the sentence that's most different from already selected ones
                max_diff = -1
                max_idx = -1
                selected_embeddings = embeddings[selected_indices]
                
                for i in range(len(sentences)):
                    if i not in selected_indices:
                        diff = torch.mean(torch.tensor([
                            torch.dist(torch.tensor(embeddings[i]), torch.tensor(sel_emb))
                            for sel_emb in selected_embeddings
                        ]))
                        if diff > max_diff:
                            max_diff = diff
                            max_idx = i
                
                selected_indices.append(max_idx)
        
        # Generate questions for selected sentences
        for idx in selected_indices:
            try:
                qa_input = sentences[idx]
                # Modified to work with text2text-generation pipeline
                generated = self.qa_pipeline(
                    f"generate question: {qa_input}",
                    max_length=64,
                    num_return_sequences=1
                )
                if generated and len(generated) > 0:
                    questions.append({
                        'question': generated[0]['generated_text'],
                        'answer': qa_input
                    })
            except Exception as e:
                st.error(f"Error generating question: {str(e)}")
                continue
                
        return questions
    
    def export_to_word(self, questions):
        doc = Document()
        doc.add_heading('Generated Exam Questions', 0)
        
        # Add questions
        doc.add_heading('Questions:', level=1)
        for i, q in enumerate(questions, 1):
            doc.add_paragraph(f"{i}. {q['question']}")
            
        # Add answer key
        doc.add_page_break()
        doc.add_heading('Answer Key:', level=1)
        for i, q in enumerate(questions, 1):
            doc.add_paragraph(f"{i}. {q['answer']}")
            
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes

def main():
    # Initialize NLTK data
    download_nltk_data()

    # Title section with gradient background
    st.markdown('<div class="title-container"><h1>📚 PDF Exam Generator</h1><p>Transform your PDF documents into professional exam questions instantly!</p></div>', unsafe_allow_html=True)

    # Quick start guide
    with st.expander("ℹ️ Quick Start Guide", expanded=not st.session_state.initialized):
        st.markdown("""
        <div class="instructions">
        <h4>How to use this tool:</h4>
        
        1. 📤 Upload your PDF document using the upload button below
        2. 🔢 Choose how many questions you want (1-10)
        3. 🎯 Click "Generate Questions"
        4. 📝 Review your questions and answers
        5. 💾 Download the complete exam as a Word document
        
        <b>Note:</b> For best results, use PDFs with clear text (not scanned documents).
        </div>
        """, unsafe_allow_html=True)
    
    # Create two columns for layout
    col1, col2 = st.columns([2, 1])

    with col1:
        st.markdown("### 📄 Upload Your Document")
        uploaded_file = st.file_uploader("Choose a PDF file", type="pdf", help="Upload a text-based PDF file")

    with col2:
        st.markdown("### ⚙️ Settings")
        num_questions = st.slider(
            "Number of questions to generate",
            min_value=1,
            max_value=10,
            value=5,
            help="Select how many questions you want to generate"
        )

    if uploaded_file is not None:
        # Create a centered container for the generate button
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            generate_button = st.button("🎯 Generate Questions", use_container_width=True)

        if generate_button:
            with st.spinner("🔄 Processing PDF and generating questions..."):
                try:
                    generator = ExamGenerator()
                    
                    # Extract and process text
                    text = generator.extract_text_from_pdf(uploaded_file)
                    sentences = generator.preprocess_text(text)
                    
                    if not sentences:
                        st.markdown('<div class="error-message">❌ Could not extract meaningful text from the PDF. Please try another file.</div>', unsafe_allow_html=True)
                        return
                    
                    # Generate questions
                    questions = generator.generate_questions(sentences, num_questions)
                    
                    if questions:
                        st.markdown(f'<div class="success-message">✅ Successfully generated {len(questions)} questions!</div>', unsafe_allow_html=True)
                        
                        # Display questions in a nice format
                        st.markdown("### 📝 Generated Questions")
                        for i, q in enumerate(questions, 1):
                            with st.expander(f"Question {i}"):
                                st.markdown(f'<div class="question-box"><strong>Q: {q["question"]}</strong><br><br>Answer: {q["answer"]}</div>', unsafe_allow_html=True)
                        
                        # Export to Word
                        doc_bytes = generator.export_to_word(questions)
                        col1, col2, col3 = st.columns([1, 2, 1])
                        with col2:
                            st.download_button(
                                label="📥 Download as Word Document",
                                data=doc_bytes,
                                file_name="generated_exam.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                    else:
                        st.markdown('<div class="error-message">⚠️ Could not generate questions. Please try another PDF file or adjust the number of questions.</div>', unsafe_allow_html=True)
                        
                except Exception as e:
                    st.markdown(f'<div class="error-message">❌ An error occurred: {str(e)}</div>', unsafe_allow_html=True)

    # Footer
    st.markdown("""
    <div class="footer">
        Made with ❤️ by AI | Free to use | No sign-up required
    </div>
    """, unsafe_allow_html=True)
    
    # Mark as initialized
    st.session_state.initialized = True

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        st.error(f"An unexpected error occurred: {str(e)}")
        st.info("Please try refreshing the page. If the error persists, contact support.") 
